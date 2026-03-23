#!/usr/bin/env python3
"""
╔══════════════════════════════════════════════════════════════╗
║           All Creator Posts to DOCX (text only) v1.0              ║
║           github.com/[your-username]/creator-posts-to-docx        ║
╚══════════════════════════════════════════════════════════════╝

Scrapes posts from a Patreon creator page and compiles them
into a formatted Word document (.docx).

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
REQUIREMENTS
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
  Python 3.8+  →  https://python.org/downloads
  (check "Add Python to PATH" during install)

  Then run once in Command Prompt / Terminal:
    pip install requests python-docx

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
STEP 1 — GET YOUR COOKIE STRING
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
  You need to be a patron of the creator to access their posts.
  The script logs in as you by using your browser's cookies.

  1. Open Chrome/Brave and log into https://www.patreon.com
  2. Navigate to the creator's posts page
       e.g. https://www.patreon.com/c/creatorname/posts
  3. Press F12 to open DevTools
  4. Click the "Network" tab
  5. In the filter box type:  api/posts
  6. Press F5 to refresh the page
  7. Click the first request that appears in the list
  8. In the right panel click "Headers"
  9. Scroll down to "Request Headers"
  10. Find the line starting with "cookie:"
  11. Copy the ENTIRE value (it's a long string)
  12. Paste it as the value of COOKIE_STRING below

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
STEP 2 — EDIT THE CONFIG BELOW
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
  Fill in COOKIE_STRING, CREATOR_URL, and OUTPUT_FILE.
  Optionally set FILTER_TITLE to grab only certain posts.

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
STEP 3 — RUN
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
  Windows:  Open Command Prompt, cd to this folder, then:
              python creator_posts_to_docx.py

  Mac/Linux: Open Terminal, cd to this folder, then:
              python3 creator_posts_to_docx.py

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
NOTES
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
  - Cookies expire after a few days. If you get an auth error,
    repeat Step 1 to get a fresh cookie string.
  - The script only accesses posts your account can already see.
    Locked/paywalled posts above your tier will be skipped.
  - Be respectful: don't scrape faster than the default delay,
    and don't distribute authors' content without permission.

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
⚠  LEGAL DISCLAIMER — READ BEFORE USE
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
  BY DOWNLOADING, INSTALLING, OR RUNNING THIS SCRIPT IN ANY
  WAY, YOU ACKNOWLEDGE THAT YOU HAVE READ, UNDERSTOOD, AND
  AGREE TO BE BOUND BY ALL OF THE FOLLOWING TERMS. IF YOU DO
  NOT AGREE, DO NOT USE THIS SCRIPT.

  1. PERSONAL USE ONLY
     This script is provided strictly for personal,
     non-commercial use. It is intended solely to allow
     paying patrons to compile content they have already
     lawfully accessed for their own private reading
     convenience. Any other use is strictly prohibited.

  2. COPYRIGHT & INTELLECTUAL PROPERTY
     All content retrieved by this script — including but
     not limited to text, stories, chapters, images, and
     any other creative works — is the sole intellectual
     property of the original content creator and is
     protected by copyright law, including but not limited
     to the U.S. Copyright Act (17 U.S.C.), the Digital
     Millennium Copyright Act (DMCA), and equivalent
     legislation worldwide (Berne Convention, EU Copyright
     Directive, etc.).

     You may NOT, under any circumstances:
       • Distribute, share, upload, or transmit the output
         to any other person or platform
       • Publish, republish, or post the content online
       • Sell, license, or monetise the content in any form
       • Use the content for AI training or data harvesting
       • Remove or alter any copyright notices
       • Create derivative works for distribution
     Doing so may constitute civil copyright infringement
     and/or criminal copyright theft, punishable by
     significant fines and/or imprisonment.

  3. NO CIRCUMVENTION OF ACCESS CONTROLS
     This script accesses only content that your own
     Patreon account is authorised to view. It does not
     bypass, crack, or circumvent any paywall or access
     control mechanism. Using it to access content beyond
     your paid tier would violate the DMCA (17 U.S.C.
     §1201) and Patreon's Terms of Service, and is
     strictly prohibited.

  4. NO WARRANTY
     THIS SCRIPT IS PROVIDED "AS IS", WITHOUT WARRANTY OF
     ANY KIND, EXPRESS OR IMPLIED, INCLUDING BUT NOT LIMITED
     TO THE WARRANTIES OF MERCHANTABILITY, FITNESS FOR A
     PARTICULAR PURPOSE, ACCURACY, OR NON-INFRINGEMENT.
     THE ENTIRE RISK AS TO THE QUALITY AND PERFORMANCE OF
     THIS SCRIPT IS WITH YOU.

  5. LIMITATION OF LIABILITY
     TO THE MAXIMUM EXTENT PERMITTED BY APPLICABLE LAW,
     IN NO EVENT SHALL THE SCRIPT'S AUTHOR(S), CONTRIBUTORS,
     DISTRIBUTORS, OR ANY ASSOCIATED PARTIES BE LIABLE FOR
     ANY DIRECT, INDIRECT, INCIDENTAL, SPECIAL, EXEMPLARY,
     CONSEQUENTIAL, OR PUNITIVE DAMAGES WHATSOEVER,
     INCLUDING BUT NOT LIMITED TO:
       • Loss of data, revenue, profits, or goodwill
       • Legal fees, fines, or penalties incurred by you
       • Claims brought against you by third parties
       • Any copyright infringement you commit
       • Any violation of Patreon's Terms of Service
       • Any damage to your computer or systems
     ARISING OUT OF OR IN CONNECTION WITH YOUR USE OR
     INABILITY TO USE THIS SCRIPT, EVEN IF ADVISED OF THE
     POSSIBILITY OF SUCH DAMAGES.

  6. INDEMNIFICATION
     By using this script you agree to fully indemnify,
     defend, and hold harmless the script's author(s),
     contributors, and distributors from and against any
     and all claims, damages, losses, liabilities, costs,
     and expenses (including reasonable legal fees) arising
     out of or relating to your use of this script, your
     violation of these terms, or your violation of any
     third-party rights, including copyright.

  7. USER RESPONSIBILITY
     You, the end user, bear sole and complete
     responsibility for:
       • How you use this script and its output
       • Ensuring your use complies with all applicable
         local, national, and international laws
       • Ensuring your use complies with Patreon's ToS
       • Any consequences, legal or otherwise, of your use
     Ignorance of these terms or applicable law is not a
     defence.

  8. SUPPORT THE CREATORS
     This tool exists to enhance your reading experience,
     not to undermine the creators whose work you enjoy.
     Please continue supporting them directly on Patreon.
     If you can no longer afford a subscription, cancel it
     — do not use this tool as a substitute for payment.
"""

import requests
import time
import json
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime


# ══════════════════════════════════════════════════════════════
# CONFIG — fill these in before running
# ══════════════════════════════════════════════════════════════

# Your Patreon cookie string (see STEP 1 above)
# Example: "session_id=abc123; patreon_device_id=xyz..."
COOKIE_STRING = "PASTE_YOUR_COOKIE_STRING_HERE"

# Full URL of the creator's Patreon page
# Example: "https://www.patreon.com/c/creatorname/"
#      or: "https://www.patreon.com/creatorname"
CREATOR_URL = "https://www.patreon.com/c/CREATOR_NAME/"

# Where to save the output file (full path recommended)
# Windows example: r"C:\Users\YourName\Downloads\book.docx"
# Mac/Linux example: "/home/yourname/Downloads/book.docx"
OUTPUT_FILE = "book.docx"

# Book title shown on the cover page
BOOK_TITLE = "My Patreon Book"

# Optional: only include posts whose title contains this text (case-insensitive)
# Set to "" to include ALL posts from the creator
# Example: "Chapter" — grabs any post with "Chapter" in the title
# Example: "My Series" — grabs only posts matching that series name
FILTER_TITLE = ""

# Sort order: "asc" = oldest post first (recommended for books)
#             "desc" = newest post first
SORT_ORDER = "asc"

# Seconds to wait between fetching each post — please don't lower below 0.5
POST_DELAY = 0.8

# ══════════════════════════════════════════════════════════════
# END OF CONFIG — no need to edit anything below this line
# ══════════════════════════════════════════════════════════════


def make_session():
    s = requests.Session()
    s.headers.update({
        "User-Agent": (
            "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
            "AppleWebKit/537.36 (KHTML, like Gecko) "
            "Chrome/146.0.0.0 Safari/537.36"
        ),
        "Accept": "application/vnd.api+json",
        "Accept-Language": "en-US,en;q=0.9",
        "Referer": CREATOR_URL,
        "sec-fetch-dest": "empty",
        "sec-fetch-mode": "cors",
        "sec-fetch-site": "same-origin",
        "dnt": "1",
        "cookie": COOKIE_STRING,
    })
    return s


def get_campaign_id(session):
    """Look up the numeric campaign ID for the creator."""
    # Handle both /creatorname and /c/creatorname URL formats
    creator_name = CREATOR_URL.rstrip("/").split("/")[-1]
    resp = session.get(
        f"https://www.patreon.com/api/campaigns"
        f"?filter[vanity]={creator_name}&json-api-version=1.0"
    )
    resp.raise_for_status()
    campaigns = resp.json().get("data", [])
    if not campaigns:
        raise ValueError(
            f"Could not find a Patreon campaign for '{creator_name}'.\n"
            f"  • Check that CREATOR_URL is correct\n"
            f"  • Make sure your COOKIE_STRING is valid and not expired"
        )
    cid = campaigns[0]["id"]
    print(f"  Campaign ID: {cid}")
    return cid


def fetch_post_list(session, campaign_id):
    """Fetch every post from the campaign, then filter and sort."""
    all_posts, page = [], 1
    url = (
        f"https://www.patreon.com/api/posts"
        f"?filter[campaign_id]={campaign_id}"
        f"&sort=-published_at"
        f"&page[count]=20"
        f"&fields[post]=title,published_at"
        f"&json-api-version=1.0"
    )
    while url:
        print(f"  Page {page}...", end=" ", flush=True)
        resp = session.get(url)
        resp.raise_for_status()
        data = resp.json()
        batch = data.get("data", [])
        all_posts.extend(batch)
        print(f"{len(batch)} posts")
        url = data.get("links", {}).get("next")
        page += 1
        time.sleep(0.5)

    # Apply title filter if set
    if FILTER_TITLE:
        filtered = [
            p for p in all_posts
            if FILTER_TITLE.lower() in (p["attributes"].get("title") or "").lower()
        ]
        print(f"  Total: {len(all_posts)} posts  |  Matching '{FILTER_TITLE}': {len(filtered)}")
    else:
        filtered = all_posts
        print(f"  Total: {len(all_posts)} posts")

    # Sort by publish date
    filtered.sort(
        key=lambda p: p["attributes"].get("published_at") or "",
        reverse=(SORT_ORDER == "desc")
    )
    return filtered


def fetch_post_content(session, post_id):
    """
    Fetch a single post's full content.
    Returns (content_string, is_locked).
    Patreon stores content in 'content_json_string' (ProseMirror JSON)
    with 'content' (HTML) as a fallback.
    """
    resp = session.get(
        f"https://www.patreon.com/api/posts/{post_id}?json-api-version=1.0"
    )
    if resp.status_code == 403:
        return None, True  # locked to higher tier
    resp.raise_for_status()
    attrs = resp.json().get("data", {}).get("attributes", {})

    # Prefer structured JSON content
    cjs = attrs.get("content_json_string")
    if cjs:
        return cjs, False

    # Fall back to HTML content
    c = attrs.get("content")
    if c:
        return c, False

    return "", False


def parse_content(raw):
    """
    Convert raw post content (ProseMirror JSON or plain text) into
    a list of (style, text) tuples for building the docx.
    Styles: 'body', 'heading2', 'heading3', 'divider'
    """
    result = []

    # Try ProseMirror JSON first
    try:
        doc = json.loads(raw)
        if doc.get("type") != "doc":
            raise ValueError("Not a ProseMirror doc")

        def walk(node):
            ntype = node.get("type", "")
            children = node.get("content", [])

            if ntype == "doc":
                for child in children:
                    walk(child)

            elif ntype in ("paragraph", "blockquote"):
                parts = []
                for child in children:
                    if child.get("type") == "text":
                        parts.append(child.get("text", ""))
                    elif child.get("type") == "hardBreak":
                        parts.append("\n")
                text = "".join(parts).strip()
                for line in text.split("\n"):
                    line = line.strip()
                    if line:
                        result.append(("body", line))

            elif ntype == "heading":
                level = node.get("attrs", {}).get("level", 2)
                parts = [c.get("text", "") for c in children if c.get("type") == "text"]
                text = "".join(parts).strip()
                if text:
                    result.append(("heading2" if level <= 2 else "heading3", text))

            elif ntype in ("bulletList", "orderedList"):
                for item in children:
                    for para in item.get("content", []):
                        parts = [
                            c.get("text", "") for c in para.get("content", [])
                            if c.get("type") == "text"
                        ]
                        text = "".join(parts).strip()
                        if text:
                            result.append(("body", f"• {text}"))

            elif ntype == "horizontalRule":
                result.append(("divider", ""))

            else:
                for child in children:
                    walk(child)

        walk(doc)
        return result

    except Exception:
        pass

    # Fall back: plain text or HTML-as-text
    for line in raw.replace("<br>", "\n").split("\n"):
        line = line.strip()
        if not line:
            continue
        # Strip any remaining HTML tags crudely
        import re
        line = re.sub(r"<[^>]+>", "", line).strip()
        if line in ("* * *", "***", "---"):
            result.append(("divider", ""))
        elif line:
            result.append(("body", line))

    return result


def build_docx(chapters):
    """Compile all chapters into a formatted Word document."""
    doc = Document()

    # Page setup — US Letter, 1.25" side margins
    for sec in doc.sections:
        sec.page_width    = int(8.5 * 914400)
        sec.page_height   = int(11  * 914400)
        sec.left_margin   = Inches(1.25)
        sec.right_margin  = Inches(1.25)
        sec.top_margin    = Inches(1.0)
        sec.bottom_margin = Inches(1.0)

    # ── Disclaimer page (always first) ──
    def add_disclaimer_page(doc):
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        tr = title.add_run("⚠  LEGAL DISCLAIMER")
        tr.font.size = Pt(16)
        tr.font.bold = True
        tr.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)

        doc.add_paragraph()

        lines = [
            ("bold",   "PERSONAL USE ONLY"),
            ("body",   "This document was compiled using an automated script for the sole purpose of personal, private reading convenience. It is not for distribution, resale, or sharing in any form."),
            ("spacer", ""),
            ("bold",   "COPYRIGHT NOTICE"),
            ("body",   "All content contained in this document is the exclusive intellectual property of the original content creator. It is protected by copyright law, including the U.S. Copyright Act, the Digital Millennium Copyright Act (DMCA), and equivalent laws worldwide. Unauthorised distribution, reproduction, uploading, selling, or sharing of this material — in whole or in part — is strictly illegal and may result in civil and/or criminal penalties."),
            ("spacer", ""),
            ("bold",   "PROHIBITED ACTIONS"),
            ("body",   "You may NOT: share or forward this file to any other person; upload it to any website, forum, or file-sharing service; sell, license, or monetise the content in any way; use the content for AI training or data harvesting; or create derivative works for distribution."),
            ("spacer", ""),
            ("bold",   "NO WARRANTY & LIMITATION OF LIABILITY"),
            ("body",   "This document is provided for personal convenience only. The script used to generate it is provided \"as is\" without warranty of any kind. The script's author(s) accept no responsibility or liability whatsoever for how this document is used. By using the script and reading this document, you accepted full and sole responsibility for your actions and released the script's author(s) from all liability."),
            ("spacer", ""),
            ("bold",   "SUPPORT THE CREATOR"),
            ("body",   "If you enjoy this content, please continue supporting the author directly on Patreon. This tool is not a substitute for a paid subscription."),
        ]

        for kind, text in lines:
            if kind == "spacer":
                doc.add_paragraph()
            elif kind == "bold":
                p = doc.add_paragraph()
                r = p.add_run(text)
                r.font.bold = True
                r.font.size = Pt(11)
            else:
                p = doc.add_paragraph(text)
                p.paragraph_format.space_after = Pt(2)
                for run in p.runs:
                    run.font.size = Pt(10)

        doc.add_page_break()

    add_disclaimer_page(doc)

    # ── Cover page ──
    cover = doc.add_paragraph()
    cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cover.add_run(BOOK_TITLE)
    r.font.size = Pt(28)
    r.font.bold = True

    doc.add_paragraph()

    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    dr = date_p.add_run(f"Compiled {datetime.today().strftime('%B %d, %Y')}")
    dr.font.size = Pt(10)
    dr.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    doc.add_page_break()

    written = skipped = 0

    for i, (title, date, paras, locked) in enumerate(chapters, 1):
        # Chapter heading
        doc.add_heading(title, level=1)

        # Publish date in small gray text
        if date:
            dp = doc.add_paragraph()
            dp.add_run(f"Published: {date}").font.size = Pt(9)
            dp.runs[0].font.color.rgb = RGBColor(0x99, 0x99, 0x99)

        doc.add_paragraph()  # spacing

        if locked:
            doc.add_paragraph("[Locked — requires a higher patron tier]")
            skipped += 1
        elif not paras:
            doc.add_paragraph("[No text content found in this post]")
        else:
            written += 1
            for style, text in paras:
                if style == "heading2":
                    doc.add_heading(text, level=2)
                elif style == "heading3":
                    doc.add_heading(text, level=3)
                elif style == "divider":
                    p = doc.add_paragraph("* * *")
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                else:
                    p = doc.add_paragraph(text)
                    p.paragraph_format.space_after = Pt(6)
                    for run in p.runs:
                        run.font.size = Pt(12)

        # Page break between chapters
        if i < len(chapters):
            doc.add_page_break()

    doc.save(OUTPUT_FILE)
    return written, skipped


def validate_config():
    """Check config before starting."""
    errors = []
    if "PASTE_YOUR_COOKIE" in COOKIE_STRING or not COOKIE_STRING.strip():
        errors.append(
            "COOKIE_STRING is not set.\n"
            "  See STEP 1 in the instructions at the top of this file."
        )
    if "CREATOR_NAME" in CREATOR_URL:
        errors.append(
            "CREATOR_URL still has the placeholder value.\n"
            "  Replace CREATOR_NAME with the actual Patreon username."
        )
    if errors:
        print("\n" + "="*60)
        print("  Setup incomplete — please fix the following:")
        print("="*60)
        for e in errors:
            print(f"\n  ✗ {e}")
        print("\nOpen this script in a text editor and fill in the CONFIG section.")
        return False
    return True


def main():
    print("\n" + "="*60)
    print("  All Creator Posts to DOCX (text only) v1.0")
    print("="*60 + "\n")

    print("⚠  LEGAL REMINDER:")
    print("   This script is for PERSONAL USE ONLY.")
    print("   All scraped content is copyright of the original creator.")
    print("   Distributing, sharing, or republishing it is ILLEGAL.")
    print("   By proceeding you accept full liability for your use")
    print("   and release the script's author(s) from all liability.")
    print()

    if not validate_config():
        return

    try:
        session = make_session()

        print("→ Finding campaign...")
        cid = get_campaign_id(session)

        print("→ Fetching post list...")
        posts = fetch_post_list(session, cid)

        if not posts:
            print(
                "\nNo posts found.\n"
                "  • If FILTER_TITLE is set, try clearing it (set to \"\") to see all posts\n"
                "  • Check the spelling in FILTER_TITLE"
            )
            return

        print(f"\n→ Fetching content for {len(posts)} posts...")
        chapters = []
        for i, post in enumerate(posts, 1):
            attrs   = post["attributes"]
            title   = attrs.get("title") or f"Post {i}"
            date    = attrs.get("published_at", "")[:10]
            post_id = post["id"]

            print(f"  [{i}/{len(posts)}] {title}", end="  ", flush=True)

            try:
                raw, locked = fetch_post_content(session, post_id)
                if locked:
                    print("[locked]")
                    chapters.append((title, date, [], True))
                elif raw:
                    paras = parse_content(raw)
                    wc = sum(len(t.split()) for _, t in paras if t)
                    print(f"[{wc} words]")
                    chapters.append((title, date, paras, False))
                else:
                    print("[empty]")
                    chapters.append((title, date, [], False))
            except Exception as e:
                print(f"[error: {e}]")
                chapters.append((title, date, [], False))

            time.sleep(POST_DELAY)

        print(f"\n→ Building DOCX...")
        written, skipped = build_docx(chapters)

        print(f"\n  Posts compiled  : {written}")
        print(f"  Locked/skipped  : {skipped}")
        print(f"\n✓ Done! Saved to:\n  {OUTPUT_FILE}")

    except requests.exceptions.HTTPError as e:
        code = e.response.status_code
        if code == 401:
            print(
                "\nERROR: Authentication failed (401).\n"
                "  Your cookies have likely expired. Repeat STEP 1 to get a\n"
                "  fresh cookie string and update COOKIE_STRING."
            )
        elif code == 403:
            print(
                "\nERROR: Access denied (403).\n"
                "  The content may be locked to a higher patron tier."
            )
        else:
            print(f"\nERROR: HTTP {code}: {e}")
    except Exception as e:
        print(f"\nUnexpected error: {e}")
        raise


if __name__ == "__main__":
    main()
